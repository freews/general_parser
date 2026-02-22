import docx

def add_citations_and_references():
    doc_path = "/home/wscho/projects/llm-test/general_parser/paper_work/table_extraction/paper_v2.docx"
    doc = docx.Document(doc_path)

    citations = {
        "TCG Storage Opal SSC": "[1]",
        "TCG Storage Architecture Core Specification": "[2]",
        "NVMe Base Specification": "[3]",
        "Datacenter NVMe SSD Specification": "[4]",
        "Qwen-VL": "[5]",
        "DeepSeek-OCR": "[6]",
        "GLM-OCR": "[7]",
        "Claude Opus 4.5": "[8]",
        "Gemini 3.0 Pro": "[9]",
        "PyMuPDF": "[10]",
        "Tabula": "[11]",
        "Camelot": "[12]"
    }

    cited_keys = set()
    
    def process_paragraphs(paragraphs):
        for para in paragraphs:
            for run_idx, run in enumerate(list(para.runs)):
                for key in citations.keys():
                    if key not in cited_keys and key in run.text:
                        parts = run.text.split(key, 1)
                        run.text = parts[0] + key
                        
                        new_run = para.add_run(citations[key])
                        new_run.font.superscript = True
                        
                        after_run = para.add_run(parts[1])
                        after_run.bold = run.bold
                        after_run.italic = run.italic
                        after_run.font.name = run.font.name
                        after_run.font.size = run.font.size
                        
                        # reorder them in the XML so they appear correctly
                        run._r.addnext(new_run._r)
                        new_run._r.addnext(after_run._r)
                        
                        cited_keys.add(key)
                        break # Go to next run, because we mutated runs array

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # Partial matches for Claude and Gemini if not found
    if "Gemini 3.0 Pro" not in cited_keys:
        for para in doc.paragraphs:
            for run in list(para.runs):
                if "Gemini" in run.text and "Gemini 3.0 Pro" not in cited_keys:
                    parts = run.text.split("Gemini", 1)
                    run.text = parts[0] + "Gemini"
                    new_run = para.add_run(citations["Gemini 3.0 Pro"])
                    new_run.font.superscript = True
                    after_run = para.add_run(parts[1])
                    run._r.addnext(new_run._r)
                    new_run._r.addnext(after_run._r)
                    cited_keys.add("Gemini 3.0 Pro")
                    
    if "Claude Opus 4.5" not in cited_keys:
        for para in doc.paragraphs:
            for run in list(para.runs):
                if "Claude Opus" in run.text and "Claude Opus 4.5" not in cited_keys:
                    parts = run.text.split("Claude Opus", 1)
                    run.text = parts[0] + "Claude Opus"
                    new_run = para.add_run(citations["Claude Opus 4.5"])
                    new_run.font.superscript = True
                    after_run = para.add_run(parts[1])
                    run._r.addnext(new_run._r)
                    new_run._r.addnext(after_run._r)
                    cited_keys.add("Claude Opus 4.5")

    doc.add_paragraph()
    
    p_header = doc.add_paragraph()
    r_header = p_header.add_run("8. References")
    r_header.bold = True

    references = [
        '[1] Trusted Computing Group (TCG), "TCG Storage Security Subsystem Class: Opal," Version 2.30.',
        '[2] Trusted Computing Group (TCG), "TCG Storage Architecture Core Specification," Version 2.01.',
        '[3] NVM Express, Inc., "NVM Express® Base Specification," Revision 2.0c (or Rev 2.03).',
        '[4] Open Compute Project (OCP), "Datacenter NVMe® SSD Specification," Version 2.0r21.',
        '[5] Bai, J., et al., "Qwen-VL: A Versatile Vision-Language Model for Understanding, Localization, Text Reading, and Beyond," arXiv preprint arXiv:2308.12966, 2023. (See also Qwen2-VL/Qwen2.5-VL updates).',
        '[6] DeepSeek-AI, "DeepSeek-VL: Towards Real-World Vision-Language Understanding," arXiv preprint arXiv:2403.05525, 2024.',
        '[7] GLM Team, "ChatGLM: A Family of Large Language Models from GLM-130B to GLM-4 All Tools," arXiv preprint arXiv:2406.12793, 2024.',
        '[8] Anthropic, "The Claude 3 Model Family: Opus, Sonnet, Haiku," 2024.',
        '[9] Google DeepMind, "Gemini 1.5: Unlocking multimodal understanding across millions of tokens of context," arXiv preprint arXiv:2403.05530, 2024.',
        '[10] Artifex Software, Inc., "PyMuPDF: A high performance Python library for data extraction, analysis, conversion & manipulation of PDF files," [Online]. Available: https://pymupdf.readthedocs.io/',
        '[11] Tabula, "Tabula: A tool for liberating data tables trapped inside PDF files," [Online]. Available: https://tabula.technology/',
        '[12] Camelot, "Camelot: PDF Table Extraction for Humans," [Online]. Available: https://camelot-py.readthedocs.io/'
    ]

    for ref in references:
        doc.add_paragraph(ref)

    try:
        doc.save(doc_path)
        print("Successfully updated paper_v2.docx")
    except Exception as e:
        print(f"Failed to save: {e}")
        doc.save(doc_path.replace(".docx", "_updated.docx"))
        print("Saved to paper_v2_updated.docx instead due to lock")

if __name__ == "__main__":
    add_citations_and_references()

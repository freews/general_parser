import docx

def list_last_paragraphs():
    doc_path = "/home/wscho/projects/llm-test/general_parser/paper_work/table_extraction/paper_v2.docx"
    doc = docx.Document(doc_path)
    
    start_idx = max(0, len(doc.paragraphs) - 100)
    for i in range(start_idx, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text:
            print(f"[{i}] {text}")

if __name__ == "__main__":
    list_last_paragraphs()

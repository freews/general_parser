import docx
from docx.shared import RGBColor

def add_figure_explanations():
    doc_path = "/home/wscho/projects/llm-test/general_parser/paper_work/table_extraction/paper_v2.docx"
    doc = docx.Document(doc_path)
    
    fig3_exp = "Parsing Challenge (Sparse Structure): This table has a highly sparse structure using implicit visual alignment rather than explicit grid lines. Empty whitespace is used to group data or imply 'same as above'. Rule-based text parsers and standard OCRs misinterpret these large gaps as column delimiters, resulting in 'ghost columns' and severe data misalignment."
    fig4_exp = "Parsing Challenge (Page-Break Fragmentation): This table physically spans across multiple pages without repeating headers. Standard page-by-page vision models and PDF parsers suffer from contextual fragmentation. They process each page in isolation, leading to truncated trailing rows at the page boundary or hallucinated new column layouts because the model lacks the stitched, holistic visual context."

    found_fig3 = False
    found_fig4 = False

    # Collect target indices
    idx_fig3, idx_fig4 = -1, -1
    for i, p in enumerate(doc.paragraphs):
        text_lower = p.text.strip().lower()
        if "figure3(sparse)" in text_lower or "figure 3(sparse)" in text_lower:
            idx_fig3 = i
        elif "figure 4(table break" in text_lower or "figure4(table break" in text_lower:
            idx_fig4 = i
            
    # Process from bottom to top so indices don't shift!
    targets = []
    if idx_fig3 != -1: targets.append((idx_fig3, fig3_exp, "found_fig3"))
    if idx_fig4 != -1: targets.append((idx_fig4, fig4_exp, "found_fig4"))
    
    targets.sort(key=lambda x: x[0], reverse=True)
    
    for idx, text, flag in targets:
        # Check if already added
        if idx + 1 < len(doc.paragraphs) and "Parsing Challenge" in doc.paragraphs[idx+1].text:
            continue
            
        # insert
        if idx + 1 < len(doc.paragraphs):
            new_p = doc.paragraphs[idx+1].insert_paragraph_before(text)
        else:
            new_p = doc.add_paragraph(text)
            
        new_p.runs[0].italic = True
        new_p.runs[0].font.color.rgb = RGBColor(105, 105, 105) # Gray
        if flag == "found_fig3": found_fig3 = True
        if flag == "found_fig4": found_fig4 = True

    if found_fig3 or found_fig4:
        try:
            doc.save(doc_path)
            print("Successfully added parsing challenge explanations to Figure 3 and/or Figure 4.")
        except PermissionError:
            print("Error: The file is open. Please close paper_v2.docx and try again.")
    else:
        print("Done. No new figures to inject or already injected.")

if __name__ == "__main__":
    add_figure_explanations()

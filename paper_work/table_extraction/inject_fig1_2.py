import docx
from docx.shared import RGBColor

def add_figure12_explanations():
    doc_path = "/home/wscho/projects/llm-test/general_parser/paper_work/table_extraction/paper_v2.docx"
    doc = docx.Document(doc_path)
    
    fig1_exp = "Parsing Challenge (Complex Merged Headers): Technical specification tables frequently employ multi-level merging in their column headers to categorize parameters. Traditional parsing heavily relies on vertical alignment and grid lines mapping down to individual data cells. When headers are complexly merged, these parsers consistently fail to anchor the sub-headers to their correct respective data columns, leading to entirely disjointed tables where values are mismatched from their defining attributes."
    
    fig2_exp = "Parsing Challenge (Nested Constraints / Table-in-Table): This table contains deeply nested hierarchies or sub-tables within individual cells. Conventional extraction techniques assume a flat, uniform grid structure and cannot comprehend arbitrary nested groupings. As a result, they flatten the nested layout unnaturally, destroying the relational hierarchy of the sub-parameters and fusing independent sub-cells into garbled text strings."

    found_fig1 = False
    found_fig2 = False

    # Collect target indices
    idx_fig1, idx_fig2 = -1, -1
    for i, p in enumerate(doc.paragraphs):
        text_lower = p.text.strip().lower()
        if "figure 1(merged" in text_lower or "figure1(" in text_lower or "figure 1 (" in text_lower:
            idx_fig1 = i
        elif "figure 2(cells" in text_lower or "figure2" in text_lower or "figure 2 (" in text_lower:
            idx_fig2 = i
            
    # Process from bottom to top so indices don't shift!
    targets = []
    if idx_fig1 != -1: targets.append((idx_fig1, fig1_exp, "found_fig1"))
    if idx_fig2 != -1: targets.append((idx_fig2, fig2_exp, "found_fig2"))
    
    targets.sort(key=lambda x: x[0], reverse=True)
    
    for idx, text, flag in targets:
        # Check if already added
        if idx + 1 < len(doc.paragraphs) and "Parsing Challenge" in doc.paragraphs[idx+1].text:
            continue
            
        # insert
        if idx + 1 < len(doc.paragraphs):
            new_p = doc.paragraphs[idx+1].insert_paragraph_before(text)
        else:
            new_p = doc.add_paragraph(text)
            
        new_p.runs[0].italic = True
        new_p.runs[0].font.color.rgb = RGBColor(105, 105, 105) # Gray
        if flag == "found_fig1": found_fig1 = True
        if flag == "found_fig2": found_fig2 = True

    if found_fig1 or found_fig2:
        try:
            doc.save(doc_path)
            print("Successfully added parsing challenge explanations to Figure 1 and/or Figure 2.")
        except PermissionError:
            print("Error: The file is open. Please close paper_v2.docx and try again.")
    else:
        print("Done. No new figures to inject or already injected.")

if __name__ == "__main__":
    add_figure12_explanations()
